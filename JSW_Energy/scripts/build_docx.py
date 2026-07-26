"""Builds the JSW Energy institutional research note (.docx).
Structure follows the Equity Research Note Master Prompt v1.2 (26 sections)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import jswdata as d

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2A, 0x78, 0xD6)
INK = RGBColor(0x0B, 0x0B, 0x0B)
GREY = RGBColor(0x52, 0x51, 0x4E)
LGREY = RGBColor(0x8A, 0x88, 0x80)
RED = RGBColor(0xC0, 0x3A, 0x2B)
HDR_FILL = "1F3864"
ALT_FILL = "EEF3FA"
SUB_FILL = "DCE6F2"

doc = Document()

# ---------------------------------------------------------------- page setup
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
for m in ("top_margin", "bottom_margin"):
    setattr(sec, m, Cm(1.7))
sec.left_margin = sec.right_margin = Cm(1.8)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(9.1)
st.font.color.rgb = INK
st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.space_before = Pt(0)
st.paragraph_format.line_spacing = 1.03



# --- schema-ordered insertion (OOXML child order is enforced) -----------------
_PPR_ORDER = ["pStyle","keepNext","keepLines","pageBreakBefore","framePr",
    "widowControl","numPr","suppressLineNumbers","pBdr","shd","tabs",
    "suppressAutoHyphens","kinsoku","wordWrap","overflowPunct","topLinePunct",
    "autoSpaceDE","autoSpaceDN","bidi","adjustRightInd","snapToGrid","spacing",
    "ind","contextualSpacing","mirrorIndents","suppressOverlap","jc",
    "textDirection","textAlignment","textboxTightWrap","outlineLvl","divId",
    "cnfStyle","rPr","sectPr","pPrChange"]
_TBLPR_ORDER = ["tblStyle","tblpPr","tblOverlap","bidiVisual","tblStyleRowBandSize",
    "tblStyleColBandSize","tblW","jc","tblCellSpacing","tblInd","tblBorders","shd",
    "tblLayout","tblCellMar","tblLook","tblCaption","tblDescription"]
_TCPR_ORDER = ["cnfStyle","tcW","gridSpan","hMerge","vMerge","tcBorders","shd",
    "noWrap","tcMar","textDirection","tcFitText","vAlign","hideMark"]


def _ordered_insert(parent, el, order):
    tag = el.tag.split("}")[1]
    idx = order.index(tag)
    for i, child in enumerate(parent):
        ctag = child.tag.split("}")[1]
        if ctag in order and order.index(ctag) > idx:
            parent.insert(i, el)
            return
    parent.append(el)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _ordered_insert(tcPr, shd, _TCPR_ORDER)


def _borders(table, colour="C9D3E0"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), colour)
        borders.append(el)
    _ordered_insert(tblPr, borders, _TBLPR_ORDER)


def _set_grid(table, widths_cm):
    """Write w:tblGrid and disable autofit so renderers honour the column widths."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:tblLayout", "w:tblW"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _ordered_insert(tblPr, layout, _TBLPR_ORDER)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tblW.set(qn("w:type"), "dxa")
    _ordered_insert(tblPr, tblW, _TBLPR_ORDER)
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)


def _keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:keepNext")):
        pPr.remove(existing)
    kn = OxmlElement("w:keepNext")
    _ordered_insert(pPr, kn, _PPR_ORDER)


def para(text="", size=9.1, bold=False, italic=False, colour=INK,
         before=0, after=5, align=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = colour
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def rich(parts, size=9.1, before=0, after=5, align=None):
    """parts: list of (text, bold, italic, colour)."""
    p = doc.add_paragraph()
    for t, b, i, c in parts:
        r = p.add_run(t)
        r.font.size = Pt(size); r.bold = b; r.italic = i
        r.font.color.rgb = c
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size = Pt(12); r.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "3"); bot.set(qn("w:color"), "1F3864")
    bdr.append(bot)
    _ordered_insert(pPr, bdr, _PPR_ORDER)
    _keep_with_next(p)
    return p


def h2(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}.  {text}")
    r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    _keep_with_next(p)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.size = Pt(9.1)
        r.font.color.rgb = INK
    r = p.add_run(text); r.font.size = Pt(9.1); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.line_spacing = 1.05
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.size = Pt(9.1)
    r = p.add_run(text); r.font.size = Pt(9.1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6)
    return p


def why(text):
    p = doc.add_paragraph()
    r = p.add_run("Why this section: ")
    r.font.size = Pt(8.5); r.bold = True; r.italic = True; r.font.color.rgb = GREY
    r = p.add_run(text)
    r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


def table(headers, rows, widths=None, first_col_left=True, size=8.2,
          bold_rows=(), sub_rows=(), note=None, hdr_size=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _borders(t)
    total = 17.4
    if widths is None:
        first = total * 0.30
        rest = (total - first) / (len(headers) - 1)
        widths = [first] + [rest] * (len(headers) - 1)
    scale = total / sum(widths)
    widths = [w * scale for w in widths]
    _set_grid(t, widths)
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        c.width = Cm(widths[i])
        _shade(c, HDR_FILL)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2.5)
        p.paragraph_format.space_after = Pt(2.5)
        r = p.add_run(htxt)
        r.bold = True
        r.font.size = Pt(hdr_size or size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) \
            else WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.width = Cm(widths[i])
            if ri in sub_rows:
                _shade(c, SUB_FILL)
            elif ri % 2 == 1:
                _shade(c, ALT_FILL)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1.6)
            p.paragraph_format.space_after = Pt(1.6)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(val))
            r.font.size = Pt(size)
            r.bold = ri in bold_rows or ri in sub_rows
            r.font.color.rgb = INK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) \
                else WD_ALIGN_PARAGRAPH.RIGHT
    if note:
        para(note, size=7.4, italic=True, colour=LGREY, before=2, after=6)
    else:
        para("", size=4, after=4)
    return t


def exhibit(png, caption, width_cm=15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(png, width=Cm(width_cm))
    para(caption, size=7.4, italic=True, colour=LGREY, after=7,
         align=WD_ALIGN_PARAGRAPH.LEFT)


def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def fmt(v, dp=0, pct=False, x=False, sign=False):
    if v is None:
        return "n/a"
    s = f"{v:,.{dp}f}"
    if sign and v > 0:
        s = "+" + s
    if pct:
        s += "%"
    if x:
        s += "x"
    return s


Y = ["FY22A", "FY23A", "FY24A", "FY25A", "FY26A", "FY27E", "FY28E", "FY29E", "FY30E"]


def yoy(series):
    out = ["n/a"]
    for i in range(1, len(series)):
        out.append(f"{(series[i] / series[i-1] - 1) * 100:,.0f}%")
    return out


# =============================================================== TITLE BLOCK
tp = doc.add_paragraph()
tp.paragraph_format.space_after = Pt(0)
r = tp.add_run("JSW ENERGY LIMITED")
r.font.size = Pt(23); r.bold = True; r.font.color.rgb = NAVY

sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(7)
r = sp.add_run("A thermal balance sheet funding a renewable company")
r.font.size = Pt(13); r.italic = True; r.font.color.rgb = BLUE

bar = doc.add_paragraph()
bar.paragraph_format.space_before = Pt(0)
bar.paragraph_format.space_after = Pt(3)
pPr = bar._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"), "1F3864")
_ordered_insert(pPr, shd, _PPR_ORDER)
for txt, bold in [("NSE: JSWENERGY", True), ("  |  CMP Rs.561", False),
                  ("  |  Mkt cap Rs.98,584cr", False), ("  |  ", False),
                  ("HOLD", True), ("  |  TP Rs.625", True),
                  ("  |  Upside +11.4%", False), ("  |  July 2026", False)]:
    r = bar.add_run(txt)
    r.font.size = Pt(11); r.bold = bold
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rich([("Market data as of 17 July 2026 (multiple history to 23 July 2026). "
       "Consolidated basis throughout. Estimates are the attached analyst model, "
       "used exactly as given; FY26 results reported 30 March 2026, 1QFY27 "
       "reported July 2026. Base target price = 11.9x FY28E EV/EBITDA less "
       "FY27E-end net debt, on 1,833mn post-QIP shares — the same construction "
       "the model carries.", False, True, GREY)], size=8.3, after=10)

# =============================================================== PART A
h1("Part A — The business")

# ---- 2 Company overview
h2(2, "Company Overview")
bullet("JSW Energy generates and sells electricity in India. It runs 5,658MW of "
       "coal and lignite thermal at Ratnagiri, Vijayanagar, Nandyal, Barmer, "
       "Ind-Barath (Odisha) and KSK Mahanadi (Chhattisgarh); 1,631MW of Himalayan "
       "hydro at Karcham Wangtoo, Baspa II and Kutehr; and 6,165MW of solar, wind "
       "and hybrid renewables held under JSW Neo Energy. Total operational capacity "
       "was 13,454MW at FY26A end and 14,535MW at 30 June 2026, of which the company "
       "says 61% is renewable.")
bullet("The company is the power arm of the Sajjan Jindal group, listed since 2010. "
       "It began as a captive supplier to JSW Steel at Vijayanagar and has grown "
       "almost entirely by acquisition and tender since: Mytrah (1,753MW, FY23), "
       "Ind-Barath (700MW, FY24–25), KSK Mahanadi (1,800MW via NCLT, FY25) and the "
       "O2 Power platform (4.7GW, Rs.12,468cr EV, FY26). There is no meaningful "
       "proprietary technology or IP — this is an assets-and-contracts business.")
bullet("FY26A was the largest step-change in the company's history: revenue rose 61% "
       "to Rs.18,901cr and EBITDA 93% to Rs.10,064cr, almost entirely because KSK "
       "Mahanadi and the O2 platform consolidated for the first time. Reported PAT "
       "rose only 15% to Rs.2,239cr — and, as Section 22 sets out, that figure was "
       "flattered by a Rs.777cr net tax credit. Net debt rose Rs.25,155cr to "
       "Rs.70,081cr over the same year.")
bullet("The stock trades at 11.2x FY28E EV/EBITDA and 3.2x FY26A book. Its own "
       "12-month forward multiple has averaged 9.0x since 2010; it sits at 13.9x "
       "today, roughly +1 standard deviation. The equity is 29% of enterprise value, "
       "so the share price is geared to the exit multiple more than to the operations.")

# ---- 3 Corporate structure
h2(3, "Corporate Structure — Standalone, Subsidiaries and JVs")
para("In plain language: the listed company you buy is a minority of the business. "
     "JSW Energy Limited on a standalone basis owns three old thermal plants "
     "totalling 2,078MW and produced 18% of FY26A segment EBITDA. Everything else "
     "sits one or two levels below it. The thermal earnings are in four separate "
     "subsidiaries; the hydro is in one; and the entire renewable build — the part "
     "of the company the valuation depends on — sits inside JSW Neo Energy Limited, "
     "a wholly-owned holding company that itself holds the acquired platforms. One "
     "of the four thermal subsidiaries, KSK Mahanadi, is only 74% owned, which "
     "matters: it was the single largest EBITDA contributor in FY26A.")
exhibit("ex7_structure.png",
        "Exhibit 1  Corporate structure. Source: FY25 annual report (AOC-1 and "
        "salient-features statement), Q4FY26 and Q1FY27 filings, analyst model.",
        width_cm=17.0)

table(
    ["Entity", "Type", "%", "What sits here", "Formed / acquired",
     "FY26A revenue (Rs.cr)", "FY26A EBITDA (Rs.cr)"],
    [
        ["JSW Energy Ltd", "Listed parent (standalone)", "—",
         "Ratnagiri 1,200MW, Vijayanagar 860MW, Nandyal 18MW", "Listed 2010",
         "3,030", "1,957"],
        ["JSW Energy (Barmer) Ltd", "Subsidiary", "100%",
         "1,080MW lignite, RERC-regulated 15% RoE", "2009–13 (as RajWest)",
         "2,382", "668"],
        ["Ind-Barath Energy (Utkal) Ltd", "Subsidiary", "100%",
         "700MW coal, no PPA, fully merchant", "Acquired FY24–25",
         "1,912", "644"],
        ["JSW Thermal Energy One Ltd", "Subsidiary", "74%",
         "KSK Mahanadi 1,800MW coal", "NCLT approval Feb 2025",
         "6,071", "3,343"],
        ["JSW Hydro Energy Ltd", "Subsidiary", "100%",
         "Karcham Wangtoo 1,091MW, Baspa II 300MW", "2011 / 2006",
         "1,112", "1,077"],
        ["JSW Energy (Kutehr) Ltd", "Step-down (under JSW Hydro)", "100%",
         "240MW hydro, 35+35yr Haryana PPA", "CoD FY26", "incl. in hydro", "incl."],
        ["JSW Neo Energy Ltd", "Subsidiary (holdco)", "100%",
         "All solar, wind, hybrid, BESS and PSP", "Formed 2020",
         "4,701", "3,994"],
        ["— O2 Power platform", "Step-down", "100%",
         "4.7GW RE platform, EV Rs.12,468cr", "Acquired FY26", "incl. above", "incl."],
        ["— Mytrah Energy", "Step-down", "100%", "1,753MW solar and wind",
         "Acquired FY23", "incl. above", "incl."],
        ["Jaigad Power Transco Ltd", "JV with MSETCL", "74%",
         "Two 400kV lines, 14% regulated RoE", "2010", "not disclosed", "in Others"],
        ["Barmer Lignite Mining Co", "JV with RSMML", "51%",
         "9 MTPA captive lignite for Barmer", "2007", "not disclosed", "in Others"],
        ["JSW Power Trading Co Ltd", "Subsidiary", "100%",
         "Merchant and exchange sales desk", "2006", "not disclosed", "in Others"],
        ["South African Coal Mining", "Subsidiary", "69%",
         "Legacy overseas coal holding", "Legacy", "not disclosed", "in Others"],
        ["Others / eliminations", "—", "—", "Holdco costs, green hydrogen, "
         "solar module project (on hold)", "—", "(307)", "(642)"],
    ],
    widths=[2.8, 1.9, 1.0, 5.1, 1.9, 1.4, 1.4], size=7.2,
    note="All entities are consolidated. Entity-level revenue and EBITDA are the "
         "reported segment disclosure (EBITDA including other income), which is the "
         "finest split published; entity-level PAT is not disclosed and is not in the "
         "model — that gap is unresolved here rather than estimated. Kutehr sits "
         "inside the hydro segment and the acquired platforms inside renewables.")

para("Which entity actually earns the money? KSK Mahanadi — acquired 17 months ago, "
     "74% owned — was the largest single contributor in FY26A at Rs.3,343cr of "
     "segment EBITDA, ahead of the whole renewable portfolio at Rs.3,994cr and more "
     "than the listed standalone entity's Rs.1,957cr. Why each other entity exists: "
     "Barmer and Jaigad Transco are regulated-return vehicles built to house "
     "cost-plus assets; Ind-Barath and KSK were distressed-asset purchases bought "
     "for cash flow at a discount to replacement cost; JSW Neo exists to ring-fence "
     "renewable project debt from the thermal balance sheet and to give the group a "
     "separately financeable — and eventually separately listable — renewable "
     "platform.")
para("Which entities absorb capital without revenue? The 'Others' segment lost "
     "Rs.642cr of EBITDA in FY26A against negative Rs.307cr of revenue. It carries "
     "holdco costs, the 3,800 tpa green hydrogen project for Tata Steel, and a 1GW "
     "solar PV module plant that the model's own tracker marks 'Project on Hold' "
     "after Rs.1,600cr of planned capex. The 500MW/1,000MWh SECI BESS-1 project, "
     "with a July 2022 letter of award, is marked 'Project under Dispute' and the "
     "model explicitly assumes no impairment against it. Ind-Barath operates without "
     "a PPA at all. Name changes and reorganisations in the period: RajWest Power "
     "became JSW Energy (Barmer); KSK Mahanadi Power Company was acquired through "
     "insolvency and now sits under JSW Thermal Energy One Ltd.")

# ---- 4 Cycles
h2(4, "Revenue & EBITDA Margin — Business Cycles")
para("Segmenting the full listed history produces six cycles. Each is defined by "
     "what the company was building and how it was paying for it.")
table(
    ["Cycle", "Revenue (Rs.cr)", "EBITDA margin", "Net debt/EBITDA",
     "ROCE", "What happened"],
    [
        ["1. FY05–FY08  Captive merchant", "491 → 1,574", "44% → 74%",
         "2.9x → 1.7x", "7% → 29%",
         "Vijayanagar only, selling to JSW Steel and merchant. Peak-cycle "
         "merchant tariffs; the best margins the company has ever earned."],
        ["2. FY09–FY13  The first build", "1,591 → 8,934", "18% → 31%",
         "20.0x → 3.6x", "2% → 9%",
         "Ratnagiri (1,200MW) and Barmer (1,080MW) commissioned. Debt to "
         "Rs.10,377cr. Margin collapsed to 18% in FY09 on imported coal."],
        ["3. FY14–FY19  Stagnation", "8,705 → 9,138", "37% → 31%",
         "2.9x → 3.6x", "8% → 7%",
         "No new capacity. Merchant glut: Vijayanagar PLF fell from 95% to 50%. "
         "EBITDA peaked at Rs.4,026cr in FY16 and fell for three years."],
        ["4. FY20–FY22  Deleveraging", "8,273 → 8,167", "36% → 44%",
         "3.3x → 2.2x", "9% → 7%",
         "Cash harvest. Net debt cut to Rs.7,758cr — the cleanest balance sheet "
         "in a decade. FY22 PAT Rs.1,729cr on flat revenue."],
        ["5. FY23–FY26A  Re-levering", "10,332 → 18,901", "32% → 53%",
         "6.3x → 7.0x", "4% → 6%",
         "Mytrah, Ind-Barath, KSK Mahanadi and O2 acquired. Net debt up 3.4x to "
         "Rs.70,081cr. Peak leverage 8.6x in FY25A."],
        ["6. FY27E–FY30E  The renewable cycle", "26,535 → 39,506", "49% → 60%",
         "6.1x → 4.9x", "6% → 9%",
         "12,451MW of new capacity, 92% of the EBITDA growth from renewables. "
         "The first cycle in which ROCE is forecast to rise."],
    ],
    widths=[3.4, 2.2, 1.7, 1.9, 1.4, 6.8], size=7.4,
    note="Source: analyst model, FY05–FY26A actuals and FY27E–FY30E estimates. "
         "Margins are consolidated EBITDA/revenue. FY09 net debt/EBITDA of 20.0x "
         "reflects Ratnagiri under construction with no earnings.")
para("What the cycles teach about the current thesis: every previous JSW Energy "
     "growth cycle ended with a lower return on capital than it began with. ROCE was "
     "29% in FY08, 9% in FY13, 7% in FY22 and 6.4% in FY26A. The FY27E–FY30E cycle "
     "is the first the model forecasts to break that pattern, and even then it "
     "reaches only 8.8% — below the 10% WACC the model's own DCF uses. The second "
     "lesson is about the balance sheet: the company was at its most valuable to "
     "shareholders in FY22, when leverage was 2.2x and it was generating free cash. "
     "It has spent the four years since converting that balance sheet capacity into "
     "capacity, and the equity market has paid for the capacity but not re-rated the "
     "returns.")

exhibit("ex1_revenue_margin.png",
        "Exhibit 2  Revenue and EBITDA margin, FY22A–FY30E, Rs. crore. Two panels "
        "share the x-axis rather than a dual y-scale. Source: analyst model.")

# ---- 5 Segments
h2(5, "Business Segments")
para("The company reports four segments under Ind AS 108 — thermal, hydro, "
     "renewables and others — and within thermal it discloses the four legal "
     "entities separately, which is unusually granular and lets each plant be mapped "
     "back to the entity in Section 3.")
table(
    ["Revenue by segment (Rs.cr)"] + d.SR_YEARS,
    [[k] + [f"{v:,.0f}" for v in vals] for k, vals in d.sr.items()],
    widths=[5.4] + [2.0] * 6, size=7.8,
    bold_rows=(4, 8), sub_rows=(8,),
    note="Source: analyst model, reported segment disclosure. The model does not "
         "carry segment revenue beyond FY28E — FY29E and FY30E segment revenue is "
         "therefore not shown rather than extrapolated. Segment EBITDA is available "
         "to FY30E and is in Exhibit 3 below.")
table(
    ["Revenue growth, yoy"] + d.SR_YEARS[1:],
    [[k] + [f"{(vals[i] / vals[i-1] - 1) * 100:,.0f}%"
            if vals[i - 1] > 0 else "nm" for i in range(1, 6)]
     for k, vals in d.sr.items() if k in
     ("Total Thermal", "Hydro", "Renewables", "Consolidated revenue")],
    widths=[5.4] + [2.4] * 5, size=7.8, bold_rows=(3,), sub_rows=(3,))

exhibit("ex2_segment_mix.png",
        "Exhibit 3  Segment EBITDA mix, FY25A–FY30E, Rs. crore. 'Others' is the "
        "balancing item between segment and consolidated EBITDA. Source: analyst model.")

bullet("thermal (5,958MW from FY27E). Sells to Rajasthan discoms (Barmer, 1,080MW on "
       "a 30-year RERC-regulated PPA at 15% RoE), to Andhra Pradesh, Tamil Nadu and "
       "Uttar Pradesh (KSK, 95% tied), to MSEDCL (300MW at Ratnagiri), to the JSW "
       "group under captive arrangements (Vijayanagar 860MW, Ratnagiri ~850MW, "
       "Nandyal 18MW) and merchant. Margin behaviour under stress is well documented: "
       "FY23 thermal EBITDA fell to Rs.1,986cr as imported coal prices spiked, "
       "because only the regulated Barmer contract is a genuine fuel pass-through. "
       "Capacity is fully built; growth is realisation-led, not volume-led, and the "
       "only volume unlock is PLF at KSK and Ind-Barath.",
       bold_lead="Thermal — 42% of FY26A revenue, 55% of FY26A EBITDA, but shrinking. ")
bullet("hydro (1,631MW). Karcham Wangtoo and Baspa II sell 88% of output on long-term "
       "PPAs running to 2046 and 2043; 12% is free power to Himachal Pradesh. Kutehr "
       "(240MW) commissioned in FY26 on a 35+35 year Haryana PPA at Rs.4.45/unit. "
       "Margins are structurally high (97% of hydro revenue converted to EBITDA in "
       "FY26A) and the only variable is hydrology — 1QFY27 hydro generation fell on "
       "weak monsoon inflows. No further hydro capacity is in the plan.",
       bold_lead="Hydro — 6% of FY26A revenue, 9% of FY26A EBITDA. ")
bullet("renewables (6,165MW at FY26A, 18,316MW by FY30E). Sells almost entirely on "
       "25-year fixed-tariff PPAs won in SECI, NTPC, SJVN, GUVNL and state tenders, "
       "plus group captive to JSW Steel. This is the only segment with material "
       "volume growth, and its margin is the reason consolidated margin rises: "
       "renewables converted 85% of revenue to EBITDA in FY26A against 49% for "
       "thermal. Growth is constrained by land, evacuation connectivity and the pace "
       "of tender awards, not by demand.",
       bold_lead="Renewables — 25% of FY26A revenue, 40% of FY26A EBITDA, 71% by FY30E. ")
bullet("holdco costs, the transmission JV, the trading desk, green hydrogen and the "
       "shelved solar module plant. It lost Rs.642cr of EBITDA in FY26A and the model "
       "assumes it turns positive from FY27E — an assumption worth Rs.816cr of the "
       "FY27E EBITDA step-up and one that is not explained in the model.",
       bold_lead="Others. ")

# ---- 6 Business model
h2(6, "Business Model & Value Chain")
para("From input to end market: for thermal, JSW Energy buys coal (imported at "
     "Ratnagiri and Vijayanagar, domestic e-auction and Shakti-linked at "
     "Ind-Barath and KSK, captive lignite at Barmer from the 51%-owned Barmer "
     "Lignite Mining JV), burns it, and sells electricity under long-term PPAs, "
     "group captive contracts or into the power exchange. For renewables and hydro "
     "there is no fuel cost at all — the plant is built once and the revenue is a "
     "fixed tariff for 25 to 40 years. Input concentration is real on the thermal "
     "side: Ratnagiri and Vijayanagar depend on imported coal and therefore carry "
     "direct rupee-dollar and seaborne-coal exposure, which is why FY23 EBITDA fell "
     "8% in a year when revenue rose 27%.")
para("Where the moat sits, in layers. There is no technology moat. The first layer "
     "is contractual: 25-to-40 year PPAs with sovereign-adjacent counterparties "
     "(SECI, NTPC, state discoms) that fix price and offtake, and a regulated 15% "
     "RoE at Barmer. The second is site and connectivity — grid evacuation capacity "
     "and land at scale in Rajasthan, Tamil Nadu, Karnataka and Gujarat take years "
     "to assemble and are the binding constraint on Indian renewable build today. "
     "The third is group: JSW Steel, JSW Cement and JSW Paints are captive offtakers "
     "for roughly 1,700MW, which is a genuinely sticky customer base and "
     "simultaneously a related-party concentration (Section 21). The fourth is cost "
     "of capital — in a reverse-auction market where every bidder buys the same "
     "panels and turbines, the cheapest funder wins the tender. JSW Energy's AA "
     "rating supports this; its 2.3x net-debt-to-equity works against it.")
para("Capital intensity is the defining feature. Gross block was Rs.88,885cr at "
     "FY26A against Rs.18,901cr of revenue — a fixed-asset turn of 0.25x, and the "
     "model has it falling to 0.24x by FY30E as more low-turn renewable and storage "
     "assets are added. At that asset turn, a business simply cannot self-fund "
     "growth: cumulative operating cash flow of Rs.70,428cr over FY27E–FY30E does "
     "not cover Rs.89,565cr of capex, let alone Rs.33,098cr of interest. Growth is "
     "funded by the balance sheet, and Section 18 sets out how.")

# ---- 7 Capacity
h2(7, "Capacity, Utilisation and Capex Plan")
exhibit("ex3_capacity.png",
        "Exhibit 4  Operational capacity by technology, FY23A–FY30E, MW. Excludes "
        "storage (875MW BESS from FY28E, 2,375MW including PSP by FY30E). "
        "Source: analyst model.")
table(
    ["Plant / block", "Location", "MW", "PLF FY26A", "PLF FY27E", "Fuel / resource",
     "Status and commissioning"],
    [
        ["Ratnagiri", "Jaigad, Maharashtra", "1,200", "72%", "74%",
         "Imported coal + domestic e-auction", "Operating since 2010–11"],
        ["Vijayanagar", "Karnataka", "860", "75%", "71%", "Imported coal, gas",
         "Operating since 2000/2009"],
        ["Nandyal", "Andhra Pradesh", "18", "68%", "90%", "Coal",
         "Operating since 2019, group captive"],
        ["Barmer", "Rajasthan", "1,080", "67%", "70%", "Captive lignite (BLMCL)",
         "Operating since 2009–13"],
        ["Ind-Barath (Utkal)", "Odisha", "700", "67%", "60%", "Shakti e-auction coal",
         "Full 700MW from FY25; no PPA"],
        ["KSK Mahanadi", "Chhattisgarh", "1,800", "78%", "75%", "FSA with Coal India",
         "Acquired FY25; units 3–6 (1,600MW) an option, not in the model"],
        ["Karcham Wangtoo", "Himachal Pradesh", "1,091", "48%", "50%", "Hydro",
         "Operating since 2011"],
        ["Baspa II", "Himachal Pradesh", "300", "48%", "50%", "Hydro",
         "Operating since 2006"],
        ["Kutehr", "Himachal Pradesh", "240", "48%", "50%", "Hydro",
         "CoD FY26 against an original target of Sep 2024 — an 18-month slip"],
        ["Solar portfolio", "Multi-state", "2,058", "21%", "25%", "Solar",
         "3,978MW by FY27E, 6,540MW by FY30E"],
        ["Wind portfolio", "Multi-state", "3,656", "25%", "25%", "Wind",
         "4,521MW by FY27E, 6,978MW by FY30E"],
        ["Hybrid / FDRE", "Multi-state", "451", "n/a", "n/a", "Solar+wind+storage",
         "521MW by FY27E, 4,798MW by FY30E"],
        ["BESS and PSP", "Rajasthan, Maharashtra, UP", "0", "n/a", "n/a",
         "Storage", "875MW BESS from FY28E; 1,500MW Bhavali PSP in FY30E. "
         "SECI BESS-1 (500MW, LoA Jul-2022) under dispute"],
        ["Salboni, West Bengal", "West Bengal", "1,600", "—", "—", "Coal",
         "Under construction, PPA signed — NOT in the model's capacity or "
         "earnings; valued separately in the SOTP"],
    ],
    widths=[2.4, 2.4, 1.1, 1.3, 1.3, 3.0, 5.9], size=7.2,
    note="Source: analyst model (Key and Assumptions sheets). PLFs are the model's "
         "assumptions. Plant-level capex and asset turns are disclosed only for the "
         "legacy thermal and hydro fleet; project-level capex for the renewable "
         "pipeline is not in the model and is therefore not shown.")
table(
    ["Capex (Rs.cr)", "FY23A", "FY24A", "FY25A", "FY26A", "FY27E", "FY28E",
     "FY29E", "FY30E"],
    [
        ["Cash capex"] + [fmt(v) for v in d.capex[1:]],
        ["Capacity added in year (MW)", "2,046", "645", "3,625", "2,579",
         "3,155", "3,387", "2,800", "3,109"],
        ["Implied capex per MW (Rs.cr)", "2.1", "12.5", "1.9", "3.9", "6.7", "6.5",
         "7.7", "8.0"],
        ["Capital work-in-progress"] + [fmt(v) for v in d.cwip[1:]],
        ["CWIP as % of gross block"] + [f"{c / g * 100:.0f}%" for c, g
                                        in zip(d.cwip[1:], d.gfa[1:])],
    ],
    widths=[4.6] + [1.6] * 8, size=7.8,
    note="Implied capex per MW is capex in the year divided by capacity commissioned "
         "in the year and is therefore distorted by construction lead times — FY24A's "
         "Rs.12.5cr/MW reflects spend on capacity that commissioned in FY25A. "
         "Maintenance capex is not separately disclosed in the model; on a "
         "Rs.88,885cr gross block, a conventional 1.5% maintenance charge would be "
         "roughly Rs.1,300cr a year, meaning approximately 94% of FY27E capex is "
         "growth capex. That split is an inference, not a disclosure.")
para("Where expansion has slipped. Three cases are visible in the model's own "
     "project tracker and all three are dated. Kutehr hydro carried an original CoD "
     "of 1 September 2024 and commissioned in FY26 — roughly 18 months late. The "
     "125MW Hetero wind acquisition carries a CoD of 31 March 2023 in the tracker "
     "and is still listed as under construction three years later. SECI BESS-1 "
     "(500MW / 1,000MWh) has a letter of award dated 15 July 2022, is marked under "
     "dispute, and the model assumes no impairment against it.")

# ---- 8 Customers
h2(8, "Customers, Contracts and Revenue Visibility")
table(
    ["Offtaker", "Capacity tied (MW)", "Contract type", "Tenure", "Who bears fuel cost"],
    [
        ["Rajasthan state discoms (Barmer)", "1,080", "Regulated cost-plus PPA",
         "30 years", "Customer — full pass-through at 15% RoE"],
        ["Andhra Pradesh, Tamil Nadu, Uttar Pradesh (KSK)", "~1,710 (95% of 1,800)",
         "Competitively-bid PPA", "Long-term", "Largely JSW — bid tariff"],
        ["JSW group captive (Steel, Cement, Paints, Severfield, Epsilon)", "~1,730",
         "Group captive PPA", "Rolling", "Passed through at Vijayanagar; related party"],
        ["Haryana, UP, Punjab, Rajasthan via PTC (hydro)", "880",
         "Long-term hydro PPA", "To 2043 / 2046", "No fuel cost"],
        ["Haryana (Kutehr)", "240", "Fixed-tariff PPA at Rs.4.45/unit",
         "35 + 35 years", "No fuel cost"],
        ["MSEDCL (Ratnagiri)", "300", "PPA on escalable-cost basis", "25 years",
         "Partly JSW — tariff is not firm for the full term"],
        ["SECI, NTPC, SJVN, GUVNL and state RE tenders", "growing to ~18,300 by FY30E",
         "25-year fixed-tariff PPA", "25 years", "No fuel cost"],
        ["Merchant / power exchange", "~13% of 1QFY27 sales", "Spot", "None",
         "JSW — full price and volume risk"],
        ["Ind-Barath (Utkal)", "700", "No PPA signed", "—",
         "JSW — the whole plant is uncontracted"],
    ],
    widths=[5.2, 2.4, 3.2, 2.2, 4.4], size=7.4,
    note="Source: analyst model (Key sheet), FY25 annual report, Q1FY27 company "
         "commentary. Top-5 and top-10 customer concentration is not disclosed by the "
         "company and is not in the model.")
para("Revenue visibility is genuinely good and genuinely concentrated at the same "
     "time. Roughly 87% of 1QFY27 sales were contracted, and the renewable build "
     "converts uncontracted capacity into 25-year fixed-tariff revenue as it "
     "commissions — that is the strongest argument in the bull case. But three "
     "things qualify it. First, 700MW at Ind-Barath has no PPA at all and is a pure "
     "merchant bet; it earned Rs.644cr of EBITDA in FY26A at a 67% PLF and a "
     "Rs.6/unit realisation, and none of that is contracted. Second, roughly 1,730MW "
     "is sold to the promoter's own steel, cement and paints companies — reliable, "
     "but a related-party arrangement whose pricing is not independently verifiable "
     "from public disclosure. Third, customer concentration is not disclosed. The "
     "company does not publish a top-5 or top-10 revenue split, so we cannot confirm "
     "diversification and treat the absence as an analytical risk rather than as "
     "evidence of a broad base.")
para("One disclosure conflict is worth naming. The Q1FY27 commentary in the model "
     "records receivable days of 'zero versus 58 days year-on-year'. The FY26A "
     "balance sheet carries Rs.2,973cr of trade receivables on Rs.18,901cr of "
     "revenue, which is 57 days. We use the balance sheet figure throughout and "
     "treat the zero-day claim as either an overdue-only measure or a data error.")

pagebreak()

# =============================================================== PART B
h1("Part B — The market")

h2(9, "Industry Overview and Industry Growth Drivers")
para("The market JSW Energy actually sells into is Indian utility-scale electricity "
     "generation. India's installed generating capacity reached about 533GW at 31 "
     "March 2026, of which 283.5GW — 53% — is non-fossil, including 274.7GW of "
     "renewables and 8.8GW of nuclear. JSW Energy's 14,535MW at June 2026 is roughly "
     "2.7% of national capacity.")
para("The structural growth drivers, separated into those already visible in "
     "reported numbers and those still prospective:")
bullet("Solar generation grew 24% and wind 11% year-on-year in Q4FY26, against total "
       "generation growth of just 3%. Renewables are taking share, not riding demand.",
       bold_lead="Visible now — renewable share gain. ")
bullet("India's stated target is 500GW of non-fossil capacity by 2030, requiring "
       "roughly 54GW of additions a year from a 283.5GW base. The tender pipeline "
       "from SECI, NTPC, SJVN and the state utilities is the mechanism, and JSW "
       "Energy's forward capacity is won through it.",
       bold_lead="Visible now — the policy mandate. ")
bullet("Firm and dispatchable renewable tenders (FDRE, hybrid, RTC) and standalone "
       "storage now carry materially better tariffs than plain solar, because the "
       "grid needs shaped power. JSW's shift into hybrid (451MW at FY26A to 4,798MW "
       "by FY30E) and 3,000MW of BESS plus 3,300MW of pumped storage is a direct "
       "response.",
       bold_lead="Prospective — the storage premium. ")
bullet("Thermal PLFs across the fleet sit at 67–78%. There is headroom, but "
       "utilisation depends on discom offtake discipline and on merchant prices, "
       "both of which are cyclical and currently soft.",
       bold_lead="Prospective — thermal utilisation recovery. ")
para("Where we sit in the cycle, and the tension in the story: generation growth of "
     "3% in Q4FY26 was the slowest fourth quarter in six years while renewable "
     "capacity kept compounding. That combination — flat demand, fast supply — is "
     "the classic setup for curtailment and for merchant tariff compression. Grid "
     "integration, transmission bottlenecks and storage readiness are the "
     "acknowledged constraints on the 500GW path. For a company whose entire "
     "forecast growth is renewable capacity, an industry that adds capacity faster "
     "than it adds demand is not an unambiguous tailwind.")
para("The demand-side counterparty is dominated by state distribution companies and "
     "central intermediaries (SECI, NTPC, SJVN). Their financial health determines "
     "payment cycles and their tender calendar determines the growth rate. This is a "
     "sovereign-adjacent, policy-driven revenue base — low credit risk in practice, "
     "but with tariff-setting power concentrated in a handful of buyers.")

h2(10, "Competitive Landscape and Market Position")
para("JSW Energy competes in two distinct markets. In thermal it competes with NTPC, "
     "Adani Power, Tata Power and the state generators for PPAs and for distressed "
     "assets; the basis of competition is tariff and balance sheet, and JSW has "
     "recently competed mainly as a buyer of stressed capacity rather than a "
     "greenfield builder. In renewables it competes in reverse auctions against "
     "Adani Green, NTPC Green, Tata Power Renewables, ReNew, Greenko, Sembcorp and a "
     "long tail of financial-sponsor platforms. In a reverse auction where everyone "
     "buys the same modules and turbines from the same suppliers, the winner is "
     "whoever has the cheapest capital and the best-secured land and evacuation. "
     "That is a real but shallow moat, and it erodes if the balance sheet tightens.")
table(
    ["Peer", "EV/EBITDA", "P/E", "Comment"],
    [
        ["JSW Energy", "11.2x FY28E / 13.9x 12m fwd", "44.0x FY26A",
         "13.5GW; 2.3x net debt/equity; ROCE 6.4% FY26A"],
        ["NTPC", "~11.0x", "~12.8x",
         "~80GW+, sovereign-owned, regulated RoE base — the sector's valuation floor"],
        ["Tata Power", "~13.8x", "~33.4x",
         "Integrated: generation, distribution, module manufacturing"],
        ["Adani Power", "~24.4x", "n/a",
         "Largest listed private thermal by market cap in 2026"],
        ["Adani Green", ">17x", ">50x",
         "Pure-play renewable; the multiple JSW's renewable arm is implicitly "
         "benchmarked against in the SOTP"],
    ],
    widths=[3.0, 3.4, 2.2, 8.8], size=7.6,
    note="Peer multiples are trailing figures collected from secondary market-data "
         "providers in July 2026 and have NOT been reconciled to audited accounts or "
         "to a common estimate basis. Treat them as indicative of relative position "
         "only. Peer revenue growth, EBITDA margin, ROCE and market capitalisation on "
         "a like-for-like basis were not obtainable and are deliberately left out "
         "rather than estimated. This is a gap in the note.")
para("Operating benchmarking is where the company looks weakest. JSW Energy's ROCE "
     "was 6.4% in FY26A and averaged 5.4% over FY23A–FY26A. NTPC, on a regulated "
     "return base, and Tata Power, on an integrated model, both run materially "
     "higher. JSW's EBITDA margin of 53.2% is high, but margin is the wrong "
     "comparison in a capital-intensive utility — a 60% margin on a 0.24x asset turn "
     "produces a 14% pre-tax return on gross assets before any leverage or "
     "depreciation, which is how a business can look highly profitable on the income "
     "statement and mediocre on capital. Market share is not the operative metric "
     "here: at 2.7% of national capacity, JSW is a price-taker in every market it "
     "sells into except the regulated Barmer contract and the group captive book.")

h2(11, "Business Growth Drivers — Company-Specific")
exhibit("ex4_bridge.png",
        "Exhibit 5  EBITDA growth bridge, FY26A to FY30E, Rs. crore. "
        "Source: analyst model, segment EBITDA.")
table(
    ["Driver", "FY26A → FY30E EBITDA (Rs.cr)", "Contribution", "Status"],
    [
        ["Renewables: 6,165MW → 18,316MW of solar, wind and hybrid",
         "4,329 → 16,767   (+12,438)", "92% of the bridge",
         "Committed for ~6,700MW under construction; probable for the tendered "
         "pipeline; the FY29E–FY30E additions are optional"],
        ["Hydro: Kutehr full-year effect", "885 → 1,329   (+443)", "3%",
         "Committed — plant commissioned FY26"],
        ["Thermal: PLF and realisation, no new capacity",
         "5,485 → 5,172   (–313)", "–2%",
         "Committed capacity, but the model forecasts thermal EBITDA to fall"],
        ["Others: holdco, transmission, trading turning positive",
         "–635 → 269   (+903)", "7%",
         "Unexplained in the model — treat as optional"],
        ["Total", "10,064 → 23,536   (+13,472)", "100%", ""],
        ["Not in the model at all", "—", "—",
         "Salboni 1,600MW thermal (under construction, PPA signed); KSK units 3–6 "
         "(1,600–1,800MW option); PSP beyond 1,500MW; any further acquisition"],
    ],
    widths=[5.6, 3.6, 2.2, 6.0], size=7.4, bold_rows=(4,), sub_rows=(4,),
    note="Source: analyst model. Status labels (committed / probable / optional) are "
         "our assessment against the model's commissioning schedule, not a company "
         "disclosure.")
para("The single number that matters in this bridge: the model implies incremental "
     "renewable capacity earns about Rs.1.02cr of EBITDA per MW (Rs.12,438cr of "
     "additional EBITDA on 12,151MW of additional capacity), against Rs.0.70cr/MW "
     "for the FY26A renewable fleet as it stands. Some of that step-up is legitimate "
     "— the FY26A fleet includes assets commissioned part-way through the year, and "
     "hybrid and FDRE tariffs are genuinely higher than plain solar. But it is an "
     "assumption of improving unit economics on a rapidly growing base, in an "
     "auction market where tariffs have generally fallen, and it is the assumption "
     "most worth tracking quarter by quarter.")

pagebreak()

# =============================================================== PART C
h1("Part C — The numbers")

h2(12, "Investment Thesis")
rich([("The variant perception. ", True, False, NAVY),
      ("The market is treating JSW Energy as a renewable compounder whose EBITDA "
       "more than doubles by FY30E and which therefore deserves a growth multiple — "
       "and it is right about the EBITDA. Where we differ is on who owns it: 92% of "
       "that growth is renewable capacity funded with debt, net debt reaches "
       "Rs.75,038cr by FY27E-end against Rs.15,942cr of FY28E EBITDA, and the equity "
       "is only 29% of enterprise value. The share price is therefore a geared bet "
       "on the exit multiple, not on the operations — and the multiple is already at "
       "+1 standard deviation of its own 16-year history.", False, False, INK)])
numbered("Capacity goes from 13,454MW at FY26A to 25,905MW by FY30E and EBITDA from "
         "Rs.10,064cr to Rs.23,536cr, a 24% CAGR. This is checkable every quarter: "
         "2,579MW commissioned in FY26A and 1,081MW since April 2026, against a "
         "3,155MW requirement for FY27E. It is not a narrative — it is a "
         "construction schedule with dates.",
         bold_lead="The build is real, contracted and largely funded. ")
numbered("Renewables converted 85% of revenue to EBITDA in FY26A against 49% for "
         "thermal, so consolidated margin rises from 53.2% to 59.6% by FY30E "
         "purely on mix. Section 16 shows the mechanism in the cost lines: fuel "
         "falls from 29.5% of revenue to 25.5% while nothing else moves much. This "
         "is arithmetic, not operating improvement.",
         bold_lead="The margin expansion is mechanical, not heroic. ")
numbered("At FY27E-end the model carries Rs.75,038cr of net debt, equal to 4.7x "
         "FY28E EBITDA. Every 1.0x of EV/EBITDA is worth Rs.87 per share, or 15.5% "
         "of the current price. At the 16-year mean multiple of 9.0x the stock is "
         "worth Rs.373; at the model's 11.9x it is worth Rs.625. Nothing about the "
         "operations changes between those two numbers.",
         bold_lead="But the equity is a thin, geared claim on that EBITDA. ")
numbered("ROCE was 6.4% in FY26A and the model has it reaching 8.8% in FY30E. The "
         "same model discounts cash flows at a 10% WACC. Nine years of growth capex "
         "and the terminal return on capital is still below the cost of capital the "
         "analyst has assigned to it. That is not a reason to be short a "
         "well-contracted asset base, but it is a decisive reason not to pay a "
         "growth multiple for it.",
         bold_lead="And the returns still do not clear the cost of capital. ")

h2(13, "Key Model Assumptions")
table(
    ["Assumption", "FY26A", "FY27E", "FY28E", "FY29E", "FY30E"],
    [
        ["Total operational capacity (MW)", "13,454", "16,609", "19,996", "22,796",
         "25,905"],
        ["— of which renewables (MW)", "6,165", "9,020", "12,407", "15,207", "18,316"],
        ["Capacity added in year (MW)", "2,579", "3,155", "3,387", "2,800", "3,109"],
        ["Storage (MW)", "0", "0", "875", "875", "2,375"],
        ["Thermal PLF (KSK / Barmer / Ratnagiri)", "78 / 67 / 72%", "75 / 70 / 74%",
         "75 / 70 / 73%", "75 / 70 / 73%", "75 / 70 / 73%"],
        ["Solar / wind PLF", "21 / 25%", "25 / 25%", "25 / 25%", "25 / 25%",
         "25 / 25%"],
        ["Revenue growth — thermal", "63%", "25%", "2%", "n/d", "n/d"],
        ["Revenue growth — renewables", "93%", "52%", "44%", "n/d", "n/d"],
        ["Revenue growth — consolidated", "61%", "40%", "14%", "17%", "11%"],
        ["EBITDA margin", "53.2%", "48.8%", "52.5%", "56.9%", "59.6%"],
        ["Fuel cost as % of revenue", "29.5%", "32.3%", "30.0%", "27.2%", "25.5%"],
        ["Staff + other cost growth", "63%", "48%", "3%", "6%", "4%"],
        ["Effective tax rate", "(39.3%)", "25.0%", "25.0%", "25.0%", "25.0%"],
        ["Capex (Rs.cr)", "10,112", "21,229", "21,864", "21,656", "24,816"],
        ["Working capital days", "44.6", "38.0", "38.0", "38.0", "38.0"],
        ["Diluted shares (mn)", "1,757", "1,833", "1,881", "1,881", "1,881"],
    ],
    widths=[5.4] + [2.4] * 5, size=7.8,
    note="Source: analyst model, used as given. 'n/d' — the model does not carry "
         "segment revenue beyond FY28E. Working capital days are held flat at 38.0 "
         "from FY27E, a modelling convention rather than a forecast.")
para("What the forecast does not assume. No Salboni (1,600MW thermal, under "
     "construction with a signed West Bengal PPA). No exercise of the KSK Mahanadi "
     "units 3–6 option (1,600–1,800MW). No pumped storage beyond the 1,500MW Bhavali "
     "project in FY30E, against a company target of 40GWh of storage by 2030. No "
     "further acquisitions. No equity beyond the Rs.4,000cr QIP at Rs.525 and the "
     "Rs.1,875cr balance on the promoter warrants at Rs.525 that are already in the "
     "share count. Two of the model's inputs are worth flagging as choices rather "
     "than forecasts: working capital days pinned at 38.0 from FY27E, and minority "
     "interest falling from Rs.523cr in FY26A to Rs.24cr in FY28E despite 26% of KSK "
     "Mahanadi being owned by someone else (see Section 22).")
para("Reconciling the model to the company's own guidance: management targets 30GW "
     "of generation capacity and 40GWh of storage by 2030. The model reaches "
     "25,905MW of generation plus 2,375MW of storage — about 28.3GW, roughly 6% "
     "short of the stated target. The estimates used here are therefore slightly "
     "more conservative than company guidance, not more aggressive.")

h2(14, "Quarterly Trend")
qi = list(range(1, 9))
table(
    ["Rs. cr"] + [d.QTRS[i] for i in qi],
    [
        ["Revenue"] + [fmt(d.q_rev[i]) for i in qi],
        ["Revenue growth, yoy"] + [f"{d.q_revyoy[i]:.0f}%" for i in qi],
        ["EBITDA"] + [fmt(d.q_ebitda[i]) for i in qi],
        ["EBITDA margin"] + [f"{d.q_margin[i]:.1f}%" for i in qi],
        ["PAT after minority"] + [fmt(d.q_pat[i]) for i in qi],
    ],
    widths=[3.0] + [1.8] * 8, size=7.8,
    note="Source: analyst model, consolidated reported quarterlies. 3QFY26 PAT "
         "includes a Rs.751cr deferred-tax credit and a Rs.65cr exceptional item.")
para("What the cadence shows. Three things. First, the year-on-year revenue growth "
     "of 60–79% through FY26 was consolidation of KSK and O2, and it has now "
     "annualised out — 1QFY27 revenue grew 1.2%. From here, growth has to come from "
     "commissioning, not from acquisition accounting. Second, the margin is "
     "genuinely improving on mix: 1QFY27 EBITDA margin of 55.2% against 54.2% a year "
     "earlier, on flat revenue, is the renewable share doing its work — EBITDA rose "
     "3% while PAT fell 37%, the entire gap being higher interest (Rs.1,519cr against "
     "Rs.1,306cr) and depreciation on newly commissioned assets. Third, and most "
     "important for anyone underwriting FY27E: the run-rate does not yet get there. "
     "1QFY27 EBITDA of Rs.2,873cr annualises to Rs.11,493cr against a FY27E estimate "
     "of Rs.12,952cr, so the remaining nine months need Rs.3,360cr a quarter, 17% "
     "above the June quarter. On PAT the gap is wider: Rs.471cr in 1QFY27 against a "
     "FY27E estimate of Rs.2,823cr requires Rs.784cr a quarter for the rest of the "
     "year. The estimate is heavily back-ended on capacity that has not yet "
     "commissioned. That is not implausible — 3,155MW is scheduled — but it is the "
     "single most fragile part of the forecast.")

h2(15, "Financial Performance")
table(
    ["Rs. cr"] + Y,
    [
        ["Revenue"] + [fmt(v) for v in d.revenue],
        ["  growth, yoy"] + yoy(d.revenue),
        ["EBITDA"] + [fmt(v) for v in d.ebitda],
        ["  EBITDA margin"] + [f"{e/r*100:.1f}%" for e, r in zip(d.ebitda, d.revenue)],
        ["Depreciation"] + [fmt(v) for v in d.dep],
        ["Other income"] + [fmt(v) for v in d.othinc],
        ["Interest expense"] + [fmt(v) for v in d.interest],
        ["PBT (reported)"] + [fmt(v) for v in d.pbt],
        ["Tax"] + [fmt(v) for v in d.tax],
        ["  effective tax rate"] + [f"{v:.1f}%" for v in d.taxrate],
        ["Minority interest"] + [fmt(v) for v in d.minority],
        ["PAT after minority"] + [fmt(v) for v in d.patrep],
        ["  PAT margin"] + [f"{p/r*100:.1f}%" for p, r in zip(d.patrep, d.revenue)],
        ["Adjusted PAT"] + [fmt(v) for v in d.patadj],
        ["EPS (Rs.)"] + [f"{v:.2f}" for v in d.eps],
        ["DPS (Rs.)"] + [f"{v:.1f}" for v in d.dps],
    ],
    widths=[2.9] + [1.61] * 9, size=7.6, bold_rows=(0, 2, 11),
    note="Source: JSW Energy reported consolidated accounts for FY22A–FY26A; FY27E–"
         "FY30E are the attached analyst model's estimates, used exactly as given. "
         "EPS is on reported PAT after minority; adjusted EPS was Rs.8.26 in FY23A "
         "(Rs.120cr exceptional gain) and Rs.13.11 in FY26A (Rs.65cr exceptional "
         "loss). FY26A's effective tax rate is negative — see Section 22.")

h2(16, "Cost Structure and Unit Economics")
table(
    ["As % of revenue"] + Y,
    [
        ["Fuel cost"] + [f"{v/r*100:.1f}%" for v, r in zip(d.fuel, d.revenue)],
        ["Electricity purchased"] + [f"{v/r*100:.1f}%" for v, r
                                     in zip(d.powpur, d.revenue)],
        ["Employee cost"] + [f"{v/r*100:.1f}%" for v, r in zip(d.staff, d.revenue)],
        ["Other expenses"] + [f"{v/r*100:.1f}%" for v, r
                              in zip(d.otherexp, d.revenue)],
        ["Total operating cost"] + [f"{(a+b+c+e)/r*100:.1f}%" for a, b, c, e, r
                                    in zip(d.fuel, d.powpur, d.staff,
                                           d.otherexp, d.revenue)],
        ["EBITDA margin"] + [f"{e/r*100:.1f}%" for e, r in zip(d.ebitda, d.revenue)],
        ["Memo: fuel cost (Rs.cr)"] + [fmt(v) for v in d.fuel],
    ],
    widths=[3.4] + [1.55] * 9, size=7.6, bold_rows=(5,), sub_rows=(5,),
    note="Source: analyst model. There is no raw-material line — fuel is the "
         "equivalent. Power and fuel are reported together in the accounts.")
para("Fixed versus variable, and the operating leverage that follows. Fuel is the "
     "only genuinely variable cost, and it is variable only with thermal volume: "
     "renewables and hydro have no fuel cost at all. Employee cost and other "
     "expenses — 3.9% and 12.9% of FY26A revenue — are substantially fixed per MW of "
     "installed capacity. That gives a cost structure that is roughly 63% fixed "
     "at the FY26A mix and heading to 75% fixed by FY30E as renewables take share. "
     "The practical consequence: a 10% shortfall in renewable generation flows "
     "through to EBITDA almost one-for-one, because there is no fuel cost to save; "
     "a 10% shortfall in thermal generation costs roughly half as much, because fuel "
     "falls with it. Operating leverage is asymmetric across the portfolio and it is "
     "getting more so.")
para("This section is what supports the margin claim made in Section 12 and in the "
     "thesis. Consolidated EBITDA margin rises from 53.2% in FY26A to 59.6% in "
     "FY30E, and the cost lines show why: fuel falls from 29.5% of revenue to 25.5% "
     "while employee cost and other expenses are broadly flat as a share. The margin "
     "expansion is a mix effect, not a cost-reduction programme, and it therefore "
     "requires no execution beyond commissioning the plants. Equally, it cannot be "
     "accelerated: there is no cost lever here.")

h2(17, "Return Ratios, Working Capital and Cash Flow")
table(
    ["Rs. cr / ratio"] + Y,
    [
        ["ROE"] + [f"{v:.1f}%" for v in d.roe],
        ["ROCE (post-tax)"] + [f"{v:.1f}%" for v in d.roce],
        ["ROCE (pre-tax)"] + [f"{v:.1f}%" for v in d.roce_pre],
        ["Operating cash flow (post-tax)"] + [fmt(v) for v in d.ocf],
        ["Capex"] + [fmt(-v) for v in d.capex],
        ["Free cash flow"] + [fmt(v) for v in d.fcf],
        ["FCF after interest"] + [fmt(v) for v in d.fcf_int],
        ["Pre-tax OCF / EBITDA"] + [f"{v:.0f}%" for v in d.ocf_ebitda],
        ["Debtor days"] + [f"{db/r*365:.0f}" for db, r in zip(d.debtors, d.revenue)],
        ["Inventory days"] + [f"{iv/r*365:.0f}" for iv, r in zip(d.invent, d.revenue)],
        ["Net working capital days"] + [f"{v:.0f}" for v in d.wcdays],
    ],
    widths=[3.6] + [1.53] * 9, size=7.6, bold_rows=(5, 6),
    note="Source: analyst model. Debtor and inventory days are computed on year-end "
         "balances over that year's revenue.")
para("Is reported profit converting into cash? At the EBITDA line, yes — pre-tax "
     "operating cash flow has run at 74% to 123% of EBITDA and the model assumes "
     "101–110% through the forecast, which is normal for a utility with contracted "
     "receivables. Debtor days improved from 81 in FY25A to 57 in FY26A. At the "
     "equity line, no, and not by a small margin. Free cash flow after interest has "
     "been negative in every year since FY22A and the model forecasts a cumulative "
     "Rs.53,993cr of negative post-interest free cash flow over FY27E–FY30E. This "
     "is not a working capital problem or an earnings-quality problem; it is what "
     "a Rs.89,565cr four-year capex programme looks like on a Rs.10,064cr EBITDA "
     "base. The company is a net consumer of capital throughout the forecast and "
     "there is no year in the model in which it becomes self-funding.")
para("Are returns improving or being diluted? Both, at different points. ROCE fell "
     "from 7.1% in FY22A to 4.2% in FY25A as acquisitions landed on the balance "
     "sheet before their earnings did, then recovered to 6.4% in FY26A as KSK "
     "consolidated for a full year. The model then has it grinding up to 8.8% by "
     "FY30E. ROE tells a flatter story — 8.1% in FY25A, 7.9% in FY26A, 11.8% in "
     "FY30E — and the Du Pont decomposition explains why the improvement is thin: "
     "asset turnover barely moves (0.16x to 0.22x), PAT margin does not recover to "
     "its FY25A level until FY30E, and the ROE improvement comes substantially from "
     "leverage rising from 2.85x to 3.45x average assets over average net worth. An "
     "ROE improvement driven by gearing is not the same thing as a better business.")

h2(18, "Treasury, Balance Sheet and Capital Allocation")
para("JSW Energy is leveraged, not cash-rich, so the debt profile is the relevant "
     "disclosure.")
table(
    ["Rs. cr"] + Y,
    [
        ["Gross debt"] + [fmt(v) for v in d.totdebt],
        ["Cash and liquid investments"] + [fmt(v) for v in d.cashinv],
        ["Net debt"] + [fmt(v) for v in d.netdebt],
        ["Net debt / EBITDA"] + [f"{v:.2f}x" for v in d.nd_ebitda],
        ["Net debt / equity"] + [f"{v:.2f}x" for v in d.nd_eq],
        ["Interest cover (EBITDA/interest)"] + [f"{v:.2f}x" for v in d.intcov],
        ["Implied cost of debt"] + ["n/a"] +
        [f"{d.interest[i] / ((d.totdebt[i] + d.totdebt[i-1]) / 2) * 100:.1f}%"
         for i in range(1, 9)],
        ["Net worth"] + [fmt(v) for v in d.networth],
        ["Book value per share (Rs.)"] + [f"{v:.0f}" for v in d.bvps],
    ],
    widths=[3.9] + [1.5] * 9, size=7.6, bold_rows=(2, 3),
    note="Source: analyst model. Implied cost of debt is interest expense over "
         "average gross debt and will overstate the marginal rate where interest on "
         "assets under construction is not capitalised. Covenant terms and the "
         "maturity profile are not disclosed in the model and could not be verified — "
         "this is a gap. Credit ratings are AA (stable) from ICRA and India Ratings "
         "per company statements.")
exhibit("ex5_leverage.png",
        "Exhibit 6  Net debt and leverage, FY22A–FY30E. Source: analyst model.")
para("The deleveraging path is real but slow and entirely growth-dependent. Net debt "
     "keeps rising in absolute terms every single year of the forecast, from "
     "Rs.70,081cr to Rs.115,314cr; leverage falls from 7.0x to 4.9x only because "
     "EBITDA rises faster. Management's stated target is below 5.0x, which on these "
     "numbers is reached in FY29E at the earliest. Interest cover of 1.70x in FY26A "
     "is the number a credit analyst would look at first, and the model has it at "
     "2.05x–2.17x through the forecast — adequate but not comfortable for an AA "
     "credit.")
para("The capital allocation record, FY22A–FY26A. Over five years the company "
     "generated about Rs.25,007cr of post-tax operating cash flow. It spent "
     "Rs.31,382cr on capex and a further Rs.26,537cr on acquiring equity in "
     "subsidiaries (principally KSK Mahanadi and O2 Power), paid roughly "
     "Rs.11,760cr of interest and about Rs.1,700cr of dividends. The gap was funded "
     "with Rs.66,953cr of new debt and Rs.6,069cr of new equity (the FY25 QIP of "
     "Rs.4,944cr and Rs.1,125cr in FY26). The result of deploying roughly "
     "Rs.73,000cr of new capital: ROCE moved from 7.1% to 6.4%. That is the honest "
     "scorecard, and it is the strongest single argument against paying a growth "
     "multiple today.")
para("What management says it will do next: 30GW of generation and 40GWh of storage "
     "by 2030, funded by internal accruals, project debt and further equity. The "
     "model builds in Rs.4,000cr of QIP at Rs.525 and Rs.1,875cr of promoter warrant "
     "conversion at Rs.525 — the promoter has already paid the 25% deposit of "
     "Rs.625cr, which is a genuine alignment signal at a strike below the current "
     "Rs.561 price. Two capital-allocation items deserve flagging. First, the "
     "company continues to pay a dividend (Rs.352cr in FY26A, rising to Rs.451cr by "
     "FY28E in the model) while running Rs.13,000cr a year of negative post-interest "
     "free cash flow — Rs.1,758cr of dividends over FY27E–FY30E funded, at the "
     "margin, with debt. Second, JSW Energy holds shares in JSW Steel worth roughly "
     "Rs.8,265cr at a Rs.1,100 reference price, an idle cross-holding in a group "
     "company while the company borrows at an implied 8–9%. The model assumes 25mn "
     "of those shares are sold in FY27E at Rs.1,260 for Rs.2,887cr net of tax. We "
     "could not verify that this disposal has been announced; it is a modelled "
     "inflow, and if it does not happen the FY27E funding plan is short by that "
     "amount.")

h2(19, "Stock Price History — Price and Business Story")
exhibit("ex6_evband.png",
        "Exhibit 7  JSW Energy 12-month forward EV/EBITDA, March 2010 to 23 July "
        "2026, with the 16-year mean and standard deviation bands. Source: analyst "
        "model (daily series).")
para("A monthly closing-price series is not in the model, so the price narrative "
     "below is told through the 12-month forward EV/EBITDA multiple, which is in the "
     "model as a daily series to 23 July 2026, together with the 52-week range of "
     "Rs.427.75 to Rs.617.35. Where a phase description refers to the share price "
     "rather than the multiple it should be read as directional, not precise. That "
     "is a gap in this section.")
bullet("The multiple derated from 9–11x to 5–6x and troughed at 3.9x in March 2020. "
       "The business explains it: no new capacity between FY14 and FY19, Vijayanagar "
       "PLF falling from 95% to 50% in a merchant glut, EBITDA peaking in FY16 and "
       "falling for three years. The market was correctly pricing a shrinking "
       "thermal utility.",
       bold_lead="FY11–FY20, the lost decade. ")
bullet("From 5.8x to 22.4x in about eighteen months. JSW Neo Energy was formed, the "
       "20GW-by-2030 target was announced, and the balance sheet was at its "
       "strongest in a decade at 2.2x net debt/EBITDA with positive free cash flow. "
       "The market paid for an option on the renewable transition before any of it "
       "was built.",
       bold_lead="FY21–FY22, the renewable re-rating. ")
bullet("Back down to a 10.3x low. FY23 EBITDA fell 8% as imported coal costs spiked, "
       "and net debt/EBITDA jumped from 2.2x to 6.3x on the Mytrah acquisition. The "
       "market discovered what the transition would cost.",
       bold_lead="FY23, the reality check. ")
bullet("Back to a 20.1x peak in 2024 on the Ind-Barath and KSK Mahanadi acquisitions "
       "and a Rs.4,944cr QIP. Cheap distressed thermal capacity was, correctly, "
       "read as accretive.",
       bold_lead="FY24–FY25, the acquisition re-rating. ")
bullet("The multiple has traded between 12.0x and 14.9x and sits at 13.9x. The "
       "revealing fact is what happened underneath it: consolidated EBITDA nearly "
       "doubled in FY26A while the multiple compressed from a 16.9x average in 2024 "
       "to 13.2x in 2026. The market has taken delivery of the EBITDA and refused to "
       "keep paying the same multiple for it, which is a rational response to net "
       "debt rising by Rs.25,155cr in the same year.",
       bold_lead="FY26A to date, consolidation. ")
para("Current price against business reality: at Rs.561 the stock is 9% below its "
     "52-week high and 31% above its 52-week low, and the forward multiple is 0.55 "
     "standard deviations above a 16-year mean that includes a decade in which this "
     "was a very different, much lower-quality company. The business today is "
     "objectively better contracted than at any point in that history. It is also "
     "three times as leveraged as it was in FY22. The multiple sitting between "
     "+1sd and the mean is, on that reading, close to fair rather than obviously "
     "wrong in either direction.")

pagebreak()

h2(20, "Valuation")
para("We value JSW Energy on EV/EBITDA. The reasons are specific rather than "
     "conventional. P/E is unusable for FY26A because the effective tax rate was "
     "negative 39% — the reported Rs.12.74 of EPS is not a repeatable number. P/BV "
     "is distorted by three acquisitions in four years and by Rs.659cr of goodwill. "
     "A DCF is unusable here for a structural reason set out below. EV/EBITDA is the "
     "right lens for a leveraged, capital-intensive, long-contract asset base — it "
     "values the assets and then makes the reader confront the debt explicitly, "
     "which in this case is the whole point.")
para("Comparables. Adani Green (>17x) is the anchor the market implicitly applies to "
     "the renewable arm, and it is not a fair one for the consolidated company: "
     "Adani Green has no thermal fleet and no merchant exposure. NTPC (~11.0x) is a "
     "fair floor — a regulated-return generator with a far stronger balance sheet, "
     "and it is a genuine question why a 2.3x-geared private generator earning 6.4% "
     "ROCE should trade above it. Tata Power (~13.8x) is the closest structural "
     "comparable — an integrated Indian private utility mid-transition — and JSW "
     "trades roughly in line with it on a forward basis. The model's own sum-of-the-"
     "parts applies 10x to thermal and hydro and 12x to renewables, which is a "
     "defensible split and is the origin of the 11.9x blended base multiple.")
table(
    ["Scenario", "Multiple", "Applied to", "EBITDA (Rs.cr)", "Enterprise value",
     "Less net debt", "Equity value", "Target price", "vs CMP"],
    [
        ["Bear", "8.0x", "FY28E", "15,942", "127,535", "(75,038)", "52,497",
         "Rs.286", "(49%)"],
        ["Downside case", "10.2x", "FY28E", "15,942", "162,607", "(75,038)",
         "87,569", "Rs.478", "(15%)"],
        ["Base", "11.9x", "FY28E", "15,942", "189,709", "(75,038)", "114,671",
         "Rs.625", "+11%"],
        ["Bull", "11.9x", "FY29E", "20,212", "240,519", "(86,670)", "153,849",
         "Rs.839", "+50%"],
        ["Reversion to 16-yr mean", "9.0x", "FY28E", "15,942", "143,477",
         "(75,038)", "68,439", "Rs.373", "(33%)"],
    ],
    widths=[2.9, 1.5, 1.5, 1.9, 2.2, 1.9, 1.9, 1.8, 1.4], size=7.4,
    bold_rows=(2,), sub_rows=(2,),
    note="Net debt deducted is gross debt less cash and current investments at the "
         "end of the prior financial year — Rs.75,038cr at FY27E-end for the FY28E "
         "cases, Rs.86,670cr at FY28E-end for the FY29E case. Shares: 1,833.5mn "
         "post-QIP, as the model uses. Multiples are the exact values in the model's "
         "valuation matrix; the row labels in that sheet read 9x/10x/12x but the "
         "cells contain 8.0x, 10.2x and 11.9x, and we have used the cells. Implied "
         "EV/EBITDA at the current price is 11.16x FY28E.")
para("Treasury and the parts the base target does not credit. The base target "
     "deducts cash and current investments but does not add back the non-current "
     "investment in JSW Steel shares, worth roughly Rs.8,265cr, or Rs.36 per share "
     "after a 20% holding-company discount. Nor does it value Salboni (1,600MW under "
     "construction) or the KSK units 3–6 option. The model's separate sum-of-the-"
     "parts, which does credit all of these and values the operational portfolio "
     "asset by asset — thermal and hydro at 1.5x price-to-book, renewables at 12x "
     "EV/EBITDA on FY28 capacity — arrives at Rs.510 per share, below the current "
     "price. The two methods bracket the market: an asset-value approach says the "
     "stock is 9% expensive, a forward-multiple approach says it is 11% cheap. That "
     "spread is the honest answer, and it is why this is a hold.")
para("Why the DCF is not used. The model's discounted cash flow produces a negative "
     "equity value of Rs.376 per share. This is not a bearish signal; it is an "
     "artefact. Free cash flow to equity is negative in four of the seven forecast "
     "years because capex runs ahead of operating cash flow throughout, and with a "
     "terminal growth rate of 5% against a 10% WACC applied to a negative terminal "
     "cash flow, the terminal value is itself negative. A DCF cannot value a company "
     "whose entire value is in assets commissioning after the forecast window. We "
     "report the output for completeness and place no weight on it.")
para("Sensitivity — what the base target becomes one notch below assumption:")
bullet("target falls to Rs.522, or 7% below the current price. An 11% upside "
       "becomes a 7% downside on a 10% earnings miss.",
       bold_lead="FY28E EBITDA 10% below forecast (Rs.14,348cr): ")
bullet("EBITDA of Rs.15,180cr and a target of Rs.576, or +3%.",
       bold_lead="FY28E EBITDA margin of 50.0% instead of 52.5%: ")
bullet("target of Rs.373, or 33% below the current price.",
       bold_lead="Multiple reverting to the 16-year mean of 9.0x: ")
bullet("EV/EBITDA values 100% of KSK Mahanadi's EBITDA but the parent owns 74%. "
       "Deducting the minority share of FY28E KSK EBITDA (Rs.661cr) at 11.9x "
       "removes Rs.43 per share, taking the base target to about Rs.582 and the "
       "upside to under 4%. Our base target of Rs.625 is the model's construction, "
       "used as given; we flag that it is generous by that amount.",
       bold_lead="Adjusting for the 26% minority in KSK Mahanadi: ")
rich([("Rating and target. ", True, False, NAVY),
      ("HOLD, base target price Rs.625, +11.4% against a CMP of Rs.561, derived as "
       "11.9x FY28E EV/EBITDA less FY27E-end net debt on 1,833mn shares. The rating "
       "falls out of the scenario table rather than the narrative: the base case "
       "offers 11% over roughly two years, the downside to a mean-reversion of the "
       "multiple is 33%, and a 10% EBITDA miss alone flips the upside negative. That "
       "is not the risk-reward of a buy.", False, False, INK)], after=8)

pagebreak()

# =============================================================== PART D
h1("Part D — The judgement")

h2(21, "Corporate Governance Checks")
table(
    ["Parameter", "Status", "Comment"],
    [
        ["Promoter holding", "66.53% (Jun-26), from 69.32% (Sep-24)",
         "The decline is dilution from the FY25 QIP and the FY27 raise, not promoter "
         "selling. Promoters have subscribed Rs.625cr of warrants at Rs.525 with "
         "Rs.1,875cr still to pay — money going in, not out."],
        ["Promoter pledge", "Not verified", "Could not be confirmed from the model "
         "or from the sources available for this note. Treat as unverified; the JSW "
         "group has historically pledged and released promoter shares across group "
         "companies. This should be checked in the latest shareholding filing before "
         "acting."],
        ["Institutional holding trend",
         "FII 14.92% → 11.41%; DII 9.78% → 16.16% over 8 quarters",
         "The direction is the signal: foreign investors have been steady sellers "
         "for two years while domestic institutions have absorbed the stock and the "
         "QIP. Combined institutional holding rose from 24.7% to 27.6%."],
        ["Free float and liquidity", "33.5% free float; Rs.98,584cr market cap",
         "Large-cap, index-weight, liquid. Daily traded value not verified."],
        ["Independent directors", "Not verified",
         "Board composition was not obtainable for this note. Named executives: "
         "Sajjan Jindal (Chairman & MD), Sharad Mahendra (Joint MD & CEO since Feb "
         "2024), Pritesh Vinay (Director Finance & CFO)."],
        ["Investor concalls", "Held quarterly",
         "The company holds quarterly earnings calls with detailed segment and "
         "commissioning disclosure — better than sector norm and reflected in the "
         "granularity of Section 5."],
        ["Related party transactions", "Material and structural, quantum not verified",
         "Roughly 1,730MW is sold to JSW Steel, JSW Cement, JSW Paints, JSW "
         "Severfield and Epsilon Carbon under group captive arrangements — see "
         "Section 3. Barmer's fuel comes from a 51% JV with a Rajasthan state "
         "entity; Jaigad Transco is a 74% JV with MSETCL. The RPT quantum in the "
         "FY26 accounts was not obtainable and is a gap."],
        ["Auditor", "Not verified", "The FY25 consolidated financials carry an "
         "independent auditor's report; we could not retrieve it to confirm whether "
         "the opinion was unmodified or whether there are key audit matters or "
         "emphasis-of-matter paragraphs."],
        ["SEBI / exchange penalties", "Not verified", "No penalties are recorded in "
         "the model. This is an absence of evidence, not evidence of absence."],
        ["Dividend track record", "Rs.2.00/share paid in each of FY22A–FY26A",
         "Unbroken but static, and a 0.36% yield. Payout has fallen from 22% to 16% "
         "as earnings grew. The model raises DPS to Rs.2.40 by FY28E."],
        ["Treasury deployment", "Rs.8,265cr of JSW Steel shares held",
         "An idle cross-holding in a group company financed, at the margin, by "
         "borrowing at 8–9%. The model assumes Rs.3,150cr of it is sold in FY27E."],
        ["Credit rating", "AA (stable), ICRA and India Ratings",
         "Reaffirmed per company statements. Interest cover of 1.70x in FY26A is "
         "thin for the rating category."],
    ],
    widths=[3.4, 4.4, 9.6], size=7.4)
para("Separating red flags from structural quirks. The structural quirks: a 74% "
     "stake in the best asset, a holdco layer for renewables, a captive relationship "
     "with the promoter's steel business, and a JV with a state mining company for "
     "fuel. None of these is a governance problem — they are how Indian power assets "
     "are financed and contracted, and the JSW Steel captive book is a genuine "
     "commercial advantage. The items we would actually flag: the idle JSW Steel "
     "cross-holding while the company borrows to fund capex; a dividend maintained "
     "through four years of large negative free cash flow; and the model's "
     "assumption that minority interest all but disappears from FY28E, which if "
     "wrong transfers roughly 18% of forecast EPS to someone else. And the honest "
     "caveat: promoter pledge, auditor opinion, board composition and related-party "
     "quantum could not be verified for this note, so this section is incomplete in "
     "exactly the places where a small governance problem would hide.")

h2(22, "Contingent Liabilities, Litigation and Accounting Quality")
para("Contingent liabilities, material litigation, tax disputes and auditor "
     "qualifications are not in the financial model, and the FY25 and FY26 annual "
     "report notes could not be retrieved for this note. That is a genuine gap and "
     "we state it rather than fill it: the size of contingent liabilities relative "
     "to a Rs.30,752cr net worth is unknown to us. Anyone acting on this note should "
     "read Note 'Contingent liabilities and commitments' in the FY26 consolidated "
     "accounts before sizing a position. What follows is what is verifiable from the "
     "model and the reported numbers.")
bullet("The consolidated tax charge in FY26A was negative Rs.777cr on a positive PBT "
       "of Rs.1,974cr — an effective rate of minus 39%, driven by a Rs.946cr "
       "deferred-tax credit in 3QFY26 and a further Rs.371cr in 4QFY26. Normalising "
       "to the 25% rate the model itself uses from FY27E would put FY26A PAT after "
       "minority at roughly Rs.969cr rather than the reported Rs.2,239cr. Reported "
       "FY26A earnings growth of 15% becomes a decline of about 50% on a normalised "
       "basis. The cause is most plausibly recognition of deferred tax assets on "
       "acquired losses at KSK Mahanadi, but the model does not say so and we have "
       "not confirmed it against the accounts.",
       bold_lead="A negative tax rate carried FY26A earnings. ")
bullet("Minority interest was Rs.523cr in FY26A — 19% of pre-minority PAT, "
       "consistent with 26% of KSK Mahanadi sitting outside the group. The model "
       "then forecasts minority interest of Rs.232cr in FY27E and Rs.24cr in FY28E, "
       "while KSK continues to generate Rs.2,543cr of EBITDA in FY28E. If minority "
       "interest instead held at its FY26A share of profit, FY28E EPS would be "
       "roughly Rs.15.72 rather than Rs.19.28 — 18% lower. No buyout of the minority "
       "appears in the cash flow statement to justify the drop.",
       bold_lead="The minority-interest assumption. ")
bullet("Capital work-in-progress stood at Rs.17,465cr at FY26A, 20% of gross block, "
       "and the model has it peaking at Rs.25,603cr in FY27E. That is normal for a "
       "company mid-build, but it is also the account where cost overruns and "
       "impairment can sit unrecognised. The 500MW SECI BESS-1 project — letter of "
       "award July 2022, marked 'Project under Dispute' — carries an explicit "
       "no-impairment assumption in the model against Rs.2,800cr of planned project "
       "cost. The 1GW solar module plant is marked 'Project on Hold' against "
       "Rs.1,600cr of planned capex.",
       bold_lead="Capitalisation and impairment risk. ")
bullet("Other current liabilities rose from Rs.2,244cr in FY24A to Rs.5,439cr in "
       "FY25A and Rs.7,654cr in FY26A, and the model then holds the line flat "
       "through FY30E. A Rs.5,400cr increase in two years is large enough to warrant "
       "reading the note; holding it flat for four years while the balance sheet "
       "grows 50% is a modelling convention, not a forecast.",
       bold_lead="A balance-sheet line worth reading. ")
bullet("Receivables at 57 days in FY26A against 81 days in FY25A, and inventory at "
       "19 days against 28 — both improving, both consistent with a shift towards "
       "contracted renewable revenue. Nothing here looks inconsistent with reported "
       "growth.",
       bold_lead="Receivable and inventory ageing. ")

h2(23, "Management Team and Guidance vs Delivery")
table(
    ["Name / role", "Background", "Our assessment"],
    [
        ["Sajjan Jindal — Chairman & Managing Director",
         "Promoter; chairman of the JSW group across steel, cement, energy, "
         "infrastructure and paints",
         "The central key-man risk and the central capital-allocation authority. "
         "Group-level capital allocation determines whether JSW Energy funds its own "
         "build or competes for group capital. Positive: the promoter is putting "
         "Rs.2,500cr into warrants at Rs.525, below the market price."],
        ["Sharad Mahendra — Joint MD & CEO (since Feb 2024)",
         "Long-serving JSW Energy executive",
         "Two years into the role and running the largest build in the company's "
         "history. Operating disclosure has improved on his watch — quarterly "
         "segment and commissioning detail is above sector norm. Unproven on a "
         "downcycle."],
        ["Pritesh Vinay — Director Finance & CFO",
         "~22 years across corporate finance, fundraising, IR, M&A and equity "
         "research",
         "The right background for a company whose central task is raising and "
         "pricing Rs.40,000cr of incremental debt and Rs.5,875cr of equity. "
         "Execution on the AA rating through a tripling of debt is the test."],
        ["Board composition", "Not verified",
         "Independent-director count and committee composition were not obtainable "
         "for this note — a gap in an otherwise assessable section."],
    ],
    widths=[3.6, 5.2, 8.6], size=7.4)
para("Guidance versus delivery. Management credibility here is testable from public "
     "documents and from the analyst's own project tracker, and the record is mixed "
     "rather than bad.")
table(
    ["What was promised", "When and where", "What was delivered"],
    [
        ["20GW of generation capacity by 2030", "announced c.2021",
         "Target raised to 30GW plus 40GWh of storage in 2026. Raised, not missed — "
         "though the analyst model reaches only 28.3GW on its own numbers, 6% short "
         "of the current target."],
        ["Kutehr hydro 240MW, CoD September 2024",
         "project tracker in the model",
         "Commissioned in FY26 — approximately 18 months late."],
        ["Hetero wind 125MW, CoD 31 March 2023", "project tracker in the model",
         "Still listed as under construction in July 2026 — over three years late."],
        ["SECI BESS-1, 500MW / 1,000MWh, LoA 15 July 2022",
         "project tracker in the model",
         "Not commissioned; marked under dispute. No impairment taken."],
        ["1GW solar PV module manufacturing", "announced, capex Rs.1,600cr",
         "Marked 'Project on Hold'."],
        ["Net debt/EBITDA below 5.0x", "management guidance, 2026",
         "6.96x at FY26A. The model reaches 5.04x in FY29E and 4.90x in FY30E — "
         "three to four years out."],
        ["FY26 capacity addition", "quarterly disclosure",
         "2,579MW added in FY26A against 3,585MW in FY25A; 1,081MW added since "
         "April 2026 including 873MW in 1QFY27. Delivery on the aggregate build has "
         "been real, even where individual projects have slipped."],
    ],
    widths=[4.4, 3.6, 9.4], size=7.4,
    note="Source: the analyst model's own project tracker (Key sheet), company "
         "quarterly disclosure and press statements. Where a promise was made on an "
         "earnings call we have not been able to cite the specific call.")
para("Verdict on the promoter structure and who it suits. This is a promoter-"
     "controlled, group-embedded utility executing the largest capital programme in "
     "its history with a balance sheet that has tripled its debt in four years. The "
     "promoter is aligned in the way that matters most — subscribing equity at "
     "Rs.525 rather than selling — and disclosure quality is genuinely good. But the "
     "aggregate delivery record is 'roughly on plan in total, consistently late on "
     "individual projects', and that is exactly the profile that turns a 24% EBITDA "
     "CAGR into an 18% one. This suits investors with a three-to-five year horizon "
     "who are underwriting an asset base rather than an earnings trajectory, and who "
     "can tolerate a 30%+ drawdown if the multiple mean-reverts. It does not suit "
     "anyone underwriting the FY27E or FY28E EPS.")

h2(24, "Triggers to Watch")
numbered("consolidated EBITDA must run at Rs.3,360cr a quarter for the remaining "
         "nine months of FY27 to reach the Rs.12,952cr estimate, against Rs.2,873cr "
         "in 1QFY27. Two consecutive quarters below Rs.3,100cr confirms the FY27E "
         "estimate is 8–10% too high.",
         bold_lead="Q2FY27 results, around October 2026 — the run-rate test. ")
numbered("the model assumes Rs.6,612cr in FY27E against Rs.3,030cr in FY26A and "
         "Rs.1,101cr delivered in 1QFY27. If nine-month standalone revenue is below "
         "Rs.4,000cr, roughly Rs.1,500cr of FY27E revenue does not exist. This is "
         "the least-explained line in the model.",
         bold_lead="Standalone segment revenue through FY27. ")
numbered("Rs.4,000cr at Rs.525 is in the share count. A raise priced below Rs.525, "
         "or larger than Rs.4,000cr, cuts the FY27E EPS of Rs.15.40 and the base "
         "target with it. Confirms or denies the funding assumption.",
         bold_lead="Completion and pricing of the modelled FY27 QIP. ")
numbered("Rs.1,875cr remains payable at Rs.525 on warrants already 25% paid. "
         "Conversion confirms promoter commitment and delivers the cash; lapse "
         "forfeits Rs.625cr and leaves a funding hole. Watch the exchange filing.",
         bold_lead="Promoter warrant conversion. ")
numbered("3,155MW is scheduled for FY27E and 873MW came in 1QFY27. A quarterly "
         "commissioning run-rate below roughly 750MW confirms slippage in the driver "
         "that is 92% of the growth story.",
         bold_lead="Quarterly MW commissioned. ")
numbered("the model books Rs.2,887cr net in FY27E from selling 25mn shares at "
         "Rs.1,260. An announcement confirms the funding plan; silence through FY27 "
         "means the plan is short by that amount and net debt ends higher than "
         "Rs.78,303cr.",
         bold_lead="Disposal of the JSW Steel cross-holding. ")

h2(25, "Key Risks")
numbered("The equity is 29% of enterprise value, so each 1.0x of EV/EBITDA is "
         "Rs.87 per share — 15.5% of the price. Reversion to the 16-year mean of "
         "9.0x takes the stock to Rs.373, a 33% loss, with no change whatsoever in "
         "the operations. What would confirm it: the forward multiple breaking "
         "below 12x while net debt keeps rising. This is the largest single risk in "
         "the name and it is a valuation risk, not a business risk.",
         bold_lead="Multiple compression — the dominant risk. ")
numbered("12,151MW of renewables must be built in four years to deliver 92% of the "
         "EBITDA growth. The company's own tracker shows Kutehr 18 months late, "
         "Hetero wind three years late and SECI BESS-1 stalled in dispute. A "
         "one-year slip across the FY28–FY30 build removes roughly Rs.3,000–"
         "4,000cr of FY30E EBITDA, which at 11.9x is over Rs.200 per share of "
         "target. What would confirm it: quarterly commissioning below 750MW.",
         bold_lead="Renewable execution slippage. ")
numbered("Cumulative free cash flow after interest is negative Rs.53,993cr over "
         "FY27E–FY30E. The plan requires Rs.40,532cr of net new debt and Rs.5,875cr "
         "of new equity, on interest cover of 1.7–2.2x and an AA rating. A rating "
         "action, a closed equity window or a repricing of project debt forces capex "
         "deferral, which directly reduces the EBITDA the valuation rests on. What "
         "would confirm it: any negative rating action, or interest cover falling "
         "below 1.7x.",
         bold_lead="Funding — the plan does not self-fund in any forecast year. ")
numbered("India's total generation grew 3% in Q4FY26 while solar generation grew "
         "24%. An oversupplied grid compresses merchant tariffs and raises "
         "curtailment. JSW has 700MW at Ind-Barath with no PPA at all and roughly "
         "13% of sales merchant. What would confirm it: sustained weakness in "
         "exchange day-ahead prices, or curtailment disclosure in the quarterly deck.",
         bold_lead="Merchant and curtailment risk in an oversupplied market. ")
numbered("FY26A reported PAT of Rs.2,239cr becomes roughly Rs.969cr at a normal tax "
         "rate. The model assumes minority interest collapses from Rs.523cr to "
         "Rs.24cr by FY28E despite 26% of KSK Mahanadi being outside the group — "
         "worth 18% of FY28E EPS. What would confirm it: the quarterly minority-"
         "interest line staying near its FY26A level.",
         bold_lead="Earnings quality — the tax credit and the minority assumption. ")
numbered("Contingent liabilities, litigation, tax disputes, auditor opinion, board "
         "composition, promoter pledge and related-party quantum could not be "
         "verified for this note. Each is a place where a material negative could "
         "sit undetected. We treat this as a risk in its own right rather than "
         "assuming the absence of evidence is favourable.",
         bold_lead="Analytical risk — what we could not verify. ")
numbered("5,958MW of coal and lignite capacity with PPAs running into the 2040s "
         "faces tightening renewable purchase obligations, potential carbon pricing "
         "and eventual stranding pressure. On a 10-to-20 year horizon this is real; "
         "within the forecast window it is not the binding constraint, which is why "
         "it is last.",
         bold_lead="Long-dated — thermal asset stranding. ")
rich([("The falsification test. ", True, False, RED),
      ("We abandon the hold and move to a sell if, at FY27 year-end, consolidated "
       "EBITDA has come in below Rs.11,500cr (the 1QFY27 run-rate, implying the "
       "build has not delivered) and net debt/EBITDA is above 6.5x (implying it has "
       "not been funded either). That combination removes both legs of the case and "
       "leaves a 4.9x-geared balance sheet with no growth to grow into it. "
       "Symmetrically, we upgrade to a buy if FY27 EBITDA exceeds Rs.13,000cr, net "
       "debt/EBITDA closes below 5.5x, and the QIP is completed at or above Rs.525 — "
       "at which point the FY29E bull case of Rs.839 becomes the base case.",
       False, False, INK)], after=8)

h2(26, "Final Outlook")
rich([("HOLD  |  Base target price Rs.625  |  Upside +11.4%  |  "
       "CMP Rs.561  |  Market cap Rs.98,584cr", True, False, NAVY)], size=11, after=7)
para("JSW Energy is in the middle of the largest transformation in its twenty-year "
     "listed history, and the transformation is real. In four years it has gone from "
     "6,605MW to 13,454MW of operational capacity, from a thermal generator with a "
     "stagnant asset base to a company where renewables will supply 71% of EBITDA by "
     "FY30E. The contracts behind that shift are 25-to-40 year fixed-tariff PPAs "
     "with sovereign-adjacent counterparties. The company holds an AA rating, runs "
     "quarterly calls with better segment disclosure than most of its peers, and its "
     "promoter is subscribing equity at Rs.525 rather than selling. None of this is "
     "a bad business.")
para("The valuation anchor and what it requires. Our base target of Rs.625 is 11.9x "
     "FY28E EV/EBITDA less FY27E-end net debt — the same construction the model "
     "carries, and a multiple consistent with where Tata Power trades and with the "
     "10x/12x thermal-and-renewables split in the sum-of-the-parts. It requires only "
     "that the model's FY28E EBITDA of Rs.15,942cr is delivered and that the market "
     "keeps paying roughly what it pays today. It requires no re-rating, no "
     "acquisition and no upside from Salboni, from KSK units 3–6, or from the "
     "Rs.36-per-share JSW Steel cross-holding, none of which are in the number.")
para("Why the risk-reward is not asymmetric, which is the reason for the rating. "
     "Three structural facts. First, net debt at FY27E-end is 4.7x FY28E EBITDA, so "
     "the equity is a 29% residual claim and every 1.0x of multiple is worth 15.5% "
     "of the share price — the downside to a mean-reverting multiple is 33% against "
     "an 11% base upside. Second, a 10% miss on FY28E EBITDA alone turns the +11% "
     "into –7%, and the 1QFY27 run-rate is currently 11% below what FY27E requires. "
     "Third, after nine years of growth capex the model's own terminal ROCE of 8.8% "
     "is still below the 10% WACC it discounts at, and the model's own asset-based "
     "sum-of-the-parts values the company at Rs.510 — below the market price. The "
     "two valuation methods bracket the current price rather than pointing the same "
     "way, which is the definition of fairly valued.")
para("What caps position sizing. The unverifiable items, principally: contingent "
     "liabilities, litigation, auditor commentary, promoter pledge and related-party "
     "quantum are all outside what this note could confirm, and they sit in a "
     "promoter-controlled group structure where 26% of the best asset already "
     "belongs to someone else. Add a funding plan that needs Rs.46,407cr of new "
     "capital across four years on 1.7–2.2x interest cover, and the correct response "
     "is a position that survives a 30% drawdown, not one that requires the bull "
     "case. We would become buyers on evidence — an FY27 EBITDA above Rs.13,000cr "
     "with leverage below 5.5x — or on a price closer to Rs.480, where the "
     "downside-case multiple of 10.2x is already in the price.")

# ---------------------------------------------------------------- footer
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Strictly Confidential — for internal circulation only. This note is "
               "for informational purposes and does not constitute investment advice.")
r.font.size = Pt(6.8); r.italic = True; r.font.color.rgb = LGREY
fp.paragraph_format.space_after = Pt(0)

fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp2.add_run("JSW Energy Limited  ·  ")
r.font.size = Pt(6.8); r.font.color.rgb = LGREY
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
run_el = OxmlElement("w:r")
rPr = OxmlElement("w:rPr")
sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "15"); rPr.append(sz)
col = OxmlElement("w:color"); col.set(qn("w:val"), "8A8880"); rPr.append(col)
run_el.append(rPr)
fld.append(run_el)
fp2._p.append(fld)

_zoom = doc.settings.element.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

doc.save("JSW_Energy_Equity_Research_Note_July2026.docx")
print("saved")
